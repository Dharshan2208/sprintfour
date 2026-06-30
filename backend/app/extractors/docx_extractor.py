"""
Extractor for Office Open XML (``.docx``) documents.

Uses the ``python-docx`` library which provides a high-level API for
walking paragraphs, tables, and sections.

**Note on pages:** DOCX is a flow-layout format — pages are computed at
render time and are not encoded in the file.  We therefore collapse the
entire document into a single virtual page.
"""

from __future__ import annotations

import os

from app.core.exceptions import DocumentExtractionException
from app.domain.models.document import Document, DocumentMetadata, Line, Page, Paragraph
from app.extractors.base_extractor import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extract text from DOCX documents."""

    format_name = "docx"

    def extract(self, file_path: str, original_filename: str) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise DocumentExtractionException(
                message="DOCX extraction requires python-docx. "
                "Install it with: pip install python-docx"
            ) from exc

        try:
            docx = DocxDocument(file_path)
        except Exception as exc:
            raise DocumentExtractionException(
                message=f"Failed to open DOCX file: {exc}"
            ) from exc

        try:
            paragraphs: list[Paragraph] = []
            global_offset = 0
            para_counter = 0

            for para_idx, docx_para in enumerate(docx.paragraphs):
                text = docx_para.text
                if not text.strip():
                    continue  # skip empty paragraphs in structural runs

                para_counter += 1

                # A DOCX paragraph usually maps to one visual line unless
                # it contains explicit line breaks (<w:br/>).
                lines: list[Line] = []
                segments = text.split("\n") if "\n" in text else [text]

                for line_idx, segment in enumerate(segments):
                    if not segment:
                        continue
                    line_start = global_offset
                    global_offset += len(segment)
                    lines.append(
                        Line(
                            text=segment,
                            line_number=line_idx + 1,
                            start_offset=line_start,
                            end_offset=global_offset,
                        )
                    )
                    # Account for the newline we split on (except the last segment)
                    if line_idx < len(segments) - 1:
                        global_offset += 1

                para_start = lines[0].start_offset if lines else global_offset
                para_end = lines[-1].end_offset if lines else global_offset
                para_full_text = "\n".join(l.text for l in lines)

                paragraphs.append(
                    Paragraph(
                        text=para_full_text,
                        paragraph_number=para_counter,
                        lines=lines,
                        start_offset=para_start,
                        end_offset=para_end,
                    )
                )

                # Inter-paragraph spacing (what python-docx separates paragraphs with)
                global_offset += 2  # \n\n

            # Build the full document text
            page_text = "\n\n".join(p.text for p in paragraphs)

            return Document(
                metadata=DocumentMetadata(
                    filename=original_filename,
                    file_size=os.path.getsize(file_path),
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    extension="docx",
                ),
                pages=[
                    Page(
                        page_number=1,
                        paragraphs=paragraphs,
                        start_offset=0,
                        end_offset=len(page_text),
                        text=page_text,
                    )
                ],
                raw_text=page_text,
                page_count=1,
                character_count=len(page_text),
            )

        except DocumentExtractionException:
            raise
        except Exception as exc:
            raise DocumentExtractionException(
                message=f"Failed to extract text from DOCX: {exc}"
            ) from exc
